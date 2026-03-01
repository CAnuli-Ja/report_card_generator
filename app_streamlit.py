import streamlit as st
import pandas as pd
from docx import Document
import io
import tempfile
import os
from pathlib import Path

# Get the directory of the current script
SCRIPT_DIR = Path(__file__).parent
PRESENTATION_HTML_PATH = SCRIPT_DIR / "PRESENTATION.html"
PRESENTATION_PPTX_PATH = SCRIPT_DIR / "PRESENTATION.pptx"

# Page configuration
st.set_page_config(
    page_title="JWFCSS Report Card Generator",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
    <style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #2c3e50;
        margin-bottom: 1rem;
    }
    .feature-box {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 0.5rem 0;
    }
    .success-box {
        background-color: #d4edda;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #28a745;
        margin: 1rem 0;
    }
    .info-box {
        background-color: #d1ecf1;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #17a2b8;
        margin: 1rem 0;
    }
    </style>
""", unsafe_allow_html=True)

# Auto comment function
def auto_comment(avg):
    try:
        avg = float(avg)
        if avg >= 75:
            return "Excellent performance. The student demonstrates strong understanding and consistent effort."
        elif avg >= 60:
            return "Good performance. Continued effort and focus will yield further improvement."
        elif avg >= 50:
            return "Satisfactory performance. Greater consistency and engagement are encouraged."
        else:
            return "Performance needs improvement. Increased effort and academic support are recommended."
    except:
        return "Performance reviewed for the term."

# Sidebar navigation
st.sidebar.title("Navigation")
page = st.sidebar.radio(
    "Select a page:",
    ["Home", "Generate Reports"],
    label_visibility="collapsed"
)

if page == "Home":
    # Display interactive presentation
    st.markdown('<div class="main-header">📋 Report Card Generator</div>', unsafe_allow_html=True)
    st.markdown("### Automated Report Card Creation Made Simple")
    st.write("Transform hours of manual work into minutes!")
    
    # Add prominent presentation download buttons
    st.divider()
    st.markdown("### 📥 Download Presentation")
    
    col_html, col_pptx = st.columns(2)
    
    with col_html:
        # Load HTML presentation file
        if PRESENTATION_HTML_PATH.exists():
            with open(PRESENTATION_HTML_PATH, 'r') as f:
                presentation_html_content = f.read()
            
            st.download_button(
                label="🌐 Interactive HTML Version",
                data=presentation_html_content,
                file_name="PRESENTATION.html",
                mime="text/html",
                use_container_width=True,
                help="Web-based presentation with keyboard navigation"
            )
            st.caption("Best for: Browser viewing, interactive slides")
        else:
            st.warning("HTML presentation file not found.")
    
    with col_pptx:
        # Load PowerPoint presentation file
        if PRESENTATION_PPTX_PATH.exists():
            with open(PRESENTATION_PPTX_PATH, 'rb') as f:
                presentation_pptx_content = f.read()
            
            st.download_button(
                label="📊 PowerPoint Version (PPTX)",
                data=presentation_pptx_content,
                file_name="PRESENTATION.pptx",
                mime="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                use_container_width=True,
                help="Professional PowerPoint presentation for Microsoft Office"
            )
            st.caption("Best for: Microsoft Office, editable slides")
        else:
            st.warning("PowerPoint presentation file not found.")
    
    st.divider()
    
    # Display presentation in iframe
    presentation_html = """
    <iframe 
        src="file:///Users/christine.anuli/Downloads/report_card_generator/PRESENTATION.html" 
        style="width:100%; height:800px; border:none; border-radius:10px;" 
        allowfullscreen>
    </iframe>
    """
    
    # Alternative: Interactive presentation slides within Streamlit
    col1, col2 = st.columns([8, 2])
    
    with col2:
        presentation_choice = st.radio(
            "View Mode:",
            ["Interactive Slides", "Features Overview"],
            label_visibility="collapsed"
        )
    
    with col1:
        if presentation_choice == "Interactive Slides":
            st.info("""
            **📊 Interactive Presentation**
            
            A full 7-minute presentation is available in two formats:
            
            **Features:**
            - 8 comprehensive slides
            - Beautiful visual design
            - Perfect for stakeholder meetings
            """)
            
            # Create download buttons for both formats
            col_h, col_p = st.columns(2)
            
            with col_h:
                if PRESENTATION_HTML_PATH.exists():
                    with open(PRESENTATION_HTML_PATH, 'r') as f:
                        presentation_html_content = f.read()
                    
                    st.download_button(
                        label="📥 Download HTML Presentation",
                        data=presentation_html_content,
                        file_name="PRESENTATION.html",
                        mime="text/html",
                        use_container_width=True
                    )
                else:
                    st.warning("HTML presentation file not found.")
            
            with col_p:
                if PRESENTATION_PPTX_PATH.exists():
                    with open(PRESENTATION_PPTX_PATH, 'rb') as f:
                        presentation_pptx_content = f.read()
                    
                    st.download_button(
                        label="📥 Download PowerPoint",
                        data=presentation_pptx_content,
                        file_name="PRESENTATION.pptx",
                        mime="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                        use_container_width=True
                    )
                else:
                    st.warning("PowerPoint presentation file not found.")
            
            st.markdown("""
            **💡 Tips:**
            - **HTML**: Open in any web browser, use arrow keys to navigate
            - **PowerPoint**: Edit directly in Microsoft Office, fully customizable
            
            Both versions cover all 8 slides:
            - Problem & Solution
            - Key Features & Benefits
            - Step-by-step Usage Guide
            - Technology Stack
            - Call to Action
            """)
        else:
            st.subheader("✨ Key Features")
            col_features = st.columns(3)
            with col_features[0]:
                st.markdown("""
                **⏱️ Save Time**
                - Generate 100+ reports in seconds
                - Automate repetitive tasks
                - Focus on what matters
                """)
            with col_features[1]:
                st.markdown("""
                **✓ Reduce Errors**
                - Consistent formatting
                - No manual copying
                - Professional results
                """)
            with col_features[2]:
                st.markdown("""
                **🔒 Secure & Private**
                - Your data stays with you
                - Cloud-based processing
                - Encrypted transfers
                """)
            
            st.divider()
            
            st.subheader("📊 How It Works")
            step_col1, step_col2, step_col3, step_col4 = st.columns(4)
            with step_col1:
                st.markdown("### 1️⃣\nUpload Excel\nwith student\ndata")
            with step_col2:
                st.markdown("### 2️⃣\nSelect the\nsheet for your\nclass")
            with step_col3:
                st.markdown("### 3️⃣\nUpload your\nWord report\ntemplate")
            with step_col4:
                st.markdown("### 4️⃣\nClick Generate\nand download\nreports!")
            
            st.divider()
            
            st.success("""
            ✅ **Ready to get started?**
            
            Click "Generate Reports" in the sidebar to upload your files and create your report cards!
            
            **Or share the full presentation with stakeholders** - use the download button at the top of this page.
            """)

else:
    # Generate Reports page
    st.markdown('<div class="main-header">Generate Report Cards</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.write("Follow these steps to generate your report cards:")
    
    # Step 1: Upload Excel file
    st.subheader("Step 1: Upload Excel Gradesheet")
    excel_file = st.file_uploader(
        "Choose your Excel file (.xlsx)",
        type=["xlsx"],
        key="excel_uploader"
    )
    
    sheet_name = None
    if excel_file is not None:
        try:
            # Read Excel to get sheet names
            xls = pd.ExcelFile(excel_file)
            sheet_names = xls.sheet_names
            
            st.success(f"✓ Excel file loaded. Found {len(sheet_names)} sheet(s)")
            
            # Step 2: Select sheet
            st.subheader("Step 2: Select Sheet Name")
            sheet_name = st.selectbox(
                "Choose a sheet:",
                sheet_names,
                key="sheet_selector"
            )
            st.write(f"Selected: **{sheet_name}**")
            
        except Exception as e:
            st.error(f"Error reading Excel file: {str(e)}")
    
    # Step 3: Upload Word template
    st.subheader("Step 3: Upload Word Report Template")
    template_file = st.file_uploader(
        "Choose your Word template (.docx)",
        type=["docx"],
        key="template_uploader"
    )
    
    if template_file is not None:
        st.success("✓ Template file uploaded")
    
    # Step 4: Generate reports
    st.subheader("Step 4: Generate Report Cards")
    
    generate_button = st.button(
        "🔨 Generate Report Cards",
        use_container_width=True,
        type="primary"
    )
    
    if generate_button:
        if excel_file is None or template_file is None or sheet_name is None:
            st.error("Please complete all steps before generating reports.")
        else:
            with st.spinner("Processing reports..."):
                try:
                    # Read Excel file
                    raw = pd.read_excel(excel_file, sheet_name=sheet_name, header=None)
                    
                    # Find the row where actual data starts
                    start_row = None
                    for idx, val in enumerate(raw[0]):
                        if pd.notna(val) and str(val).strip().isdigit():
                            start_row = idx
                            break
                    
                    if start_row is None:
                        st.error("Could not find student data in Excel file. Make sure the first column contains student numbers.")
                    else:
                        # Extract data
                        df = raw.iloc[start_row:].reset_index(drop=True)
                        df.columns = df.iloc[0]
                        df = df[1:].reset_index(drop=True)
                        
                        if len(df) == 0:
                            st.error("No student data found in Excel file.")
                        else:
                            # Create temporary directory for generated files
                            with tempfile.TemporaryDirectory() as tmpdir:
                                count = 0
                                files_created = []
                                
                                for _, row in df.iterrows():
                                    try:
                                        doc = Document(io.BytesIO(template_file.read()))
                                        template_file.seek(0)  # Reset file pointer
                                        
                                        # Get student data
                                        first_name = str(row.iloc[1]).strip() if pd.notna(row.iloc[1]) else ""
                                        last_name = str(row.iloc[2]).strip() if pd.notna(row.iloc[2]) else ""
                                        full_name = f"{first_name} {last_name}".strip()
                                        
                                        if not full_name:
                                            continue
                                        
                                        # Build replacements dictionary
                                        replacements = {
                                            "{{NAME}}": full_name,
                                            "{{HFL}}": str(row.iloc[3]).strip() if pd.notna(row.iloc[3]) else "",
                                            "{{LIT}}": str(row.iloc[4]).strip() if pd.notna(row.iloc[4]) else "",
                                            "{{REL}}": str(row.iloc[5]).strip() if pd.notna(row.iloc[5]) else "",
                                            "{{ELANG}}": str(row.iloc[6]).strip() if pd.notna(row.iloc[6]) else "",
                                            "{{ELIT}}": str(row.iloc[7]).strip() if pd.notna(row.iloc[7]) else "",
                                            "{{FR}}": str(row.iloc[8]).strip() if pd.notna(row.iloc[8]) else "",
                                            "{{SPN}}": str(row.iloc[9]).strip() if pd.notna(row.iloc[9]) else "",
                                            "{{SSTUD}}": str(row.iloc[10]).strip() if pd.notna(row.iloc[10]) else "",
                                            "{{HIST}}": str(row.iloc[11]).strip() if pd.notna(row.iloc[11]) else "",
                                            "{{MATH}}": str(row.iloc[12]).strip() if pd.notna(row.iloc[12]) else "",
                                            "{{AGR}}": str(row.iloc[13]).strip() if pd.notna(row.iloc[13]) else "",
                                            "{{INTSCI}}": str(row.iloc[14]).strip() if pd.notna(row.iloc[14]) else "",
                                            "{{PE}}": str(row.iloc[15]).strip() if pd.notna(row.iloc[15]) else "",
                                            "{{IT}}": str(row.iloc[16]).strip() if pd.notna(row.iloc[16]) else "",
                                            "{{VA}}": str(row.iloc[17]).strip() if pd.notna(row.iloc[17]) else "",
                                            "{{HE}}": str(row.iloc[18]).strip() if pd.notna(row.iloc[18]) else "",
                                            "{{TOTAL}}": str(row.iloc[19]).strip() if pd.notna(row.iloc[19]) else "",
                                            "{{SAVG}}": str(row.iloc[20]).strip() if pd.notna(row.iloc[20]) else "",
                                            "{{TOTAL_SUBJECTS}}": str(row.iloc[21]).strip() if pd.notna(row.iloc[21]) else "",
                                            "{{PASSED}}": str(row.iloc[22]).strip() if pd.notna(row.iloc[22]) else "",
                                            "{{FORM_TEACHER_COMMENTS}}": str(row.iloc[23]).strip() if pd.notna(row.iloc[23]) else "",
                                            "{{PRINCIPAL_COMMENTS}}": str(row.iloc[24]).strip() if pd.notna(row.iloc[24]) else "",
                                            "{{ATA}}": str(row.iloc[25]).strip() if pd.notna(row.iloc[25]) else "",
                                            "{{RESP}}": str(row.iloc[26]).strip() if pd.notna(row.iloc[26]) else "",
                                            "{{CO_OP}}": str(row.iloc[27]).strip() if pd.notna(row.iloc[27]) else "",
                                            "{{ATC}}": str(row.iloc[28]).strip() if pd.notna(row.iloc[28]) else "",
                                            "{{LEAD}}": str(row.iloc[29]).strip() if pd.notna(row.iloc[29]) else "",
                                            "{{DEP}}": str(row.iloc[30]).strip() if pd.notna(row.iloc[30]) else "",
                                            "{{SOC}}": str(row.iloc[31]).strip() if pd.notna(row.iloc[31]) else "",
                                            "{{INIT}}": str(row.iloc[32]).strip() if pd.notna(row.iloc[32]) else "",
                                            "{{CON_MGT}}": str(row.iloc[33]).strip() if pd.notna(row.iloc[33]) else "",
                                            "{{APP}}": str(row.iloc[34]).strip() if pd.notna(row.iloc[34]) else "",
                                        }
                                        
                                        # Replace text in paragraphs
                                        for p in doc.paragraphs:
                                            for placeholder, value in replacements.items():
                                                if placeholder in p.text:
                                                    full_text = p.text
                                                    new_text = full_text.replace(placeholder, value)
                                                    for run in p.runs:
                                                        run.text = ""
                                                    if p.runs:
                                                        p.runs[0].text = new_text
                                                    else:
                                                        p.add_run(new_text)
                                        
                                        # Replace text in tables
                                        for table in doc.tables:
                                            for row_table in table.rows:
                                                for cell in row_table.cells:
                                                    for p in cell.paragraphs:
                                                        for placeholder, value in replacements.items():
                                                            if placeholder in p.text:
                                                                full_text = p.text
                                                                new_text = full_text.replace(placeholder, value)
                                                                for run in p.runs:
                                                                    run.text = ""
                                                                if p.runs:
                                                                    p.runs[0].text = new_text
                                                                else:
                                                                    p.add_run(new_text)
                                        
                                        # Save document
                                        filename = f"{full_name.replace(' ', '_')}_Report.docx"
                                        filepath = os.path.join(tmpdir, filename)
                                        doc.save(filepath)
                                        files_created.append((filename, filepath))
                                        count += 1
                                    
                                    except Exception as e:
                                        st.warning(f"Error processing {full_name}: {str(e)}")
                                
                                # Display success message
                                st.markdown(
                                    f"""
                                    <div class="success-box">
                                    <strong>✓ Success!</strong> Generated {count} report card(s)
                                    </div>
                                    """,
                                    unsafe_allow_html=True
                                )
                                
                                # Provide download options
                                if files_created:
                                    st.subheader("📥 Download Reports")
                                    
                                    # Create a zip file with all reports
                                    import zipfile
                                    zip_buffer = io.BytesIO()
                                    
                                    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                                        for filename, filepath in files_created:
                                            with open(filepath, 'rb') as f:
                                                zip_file.writestr(filename, f.read())
                                    
                                    zip_buffer.seek(0)
                                    
                                    st.download_button(
                                        label="📦 Download All Reports (ZIP)",
                                        data=zip_buffer.getvalue(),
                                        file_name=f"Report_Cards_{sheet_name}.zip",
                                        mime="application/zip",
                                        use_container_width=True
                                    )
                                    
                                    st.divider()
                                    
                                    st.write("**Individual Downloads:**")
                                    cols = st.columns(3)
                                    for idx, (filename, filepath) in enumerate(files_created):
                                        with open(filepath, 'rb') as f:
                                            cols[idx % 3].download_button(
                                                label=filename.replace("_Report.docx", ""),
                                                data=f.read(),
                                                file_name=filename,
                                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                            )
                
                except Exception as e:
                    st.error(f"Error generating reports: {str(e)}")
    
    # Placeholder section
    st.divider()
    st.subheader("📝 Available Placeholders for Your Template")
    
    with st.expander("View all placeholders"):
        placeholder_data = {
            "Student Information": {
                "{{NAME}}": "Full student name"
            },
            "Subjects": {
                "{{MATH}}": "Mathematics",
                "{{ELANG}}": "English Language",
                "{{ELIT}}": "English Literature",
                "{{FR}}": "French",
                "{{SPN}}": "Spanish",
                "{{SSTUD}}": "Social Studies",
                "{{HIST}}": "History",
                "{{HFL}}": "Health, Family & Life",
                "{{LIT}}": "Literacy",
                "{{REL}}": "Religion",
                "{{AGR}}": "Agricultural Science",
                "{{INTSCI}}": "Integrated Science",
                "{{PE}}": "Physical Education",
                "{{IT}}": "Information Technology",
                "{{VA}}": "Visual Art",
                "{{HE}}": "Home Economics"
            },
            "Statistics": {
                "{{TOTAL}}": "Total score",
                "{{SAVG}}": "Student average",
                "{{TOTAL_SUBJECTS}}": "Total number of subjects",
                "{{PASSED}}": "Number of subjects passed"
            },
            "Comments": {
                "{{FORM_TEACHER_COMMENTS}}": "Form teacher's comments",
                "{{PRINCIPAL_COMMENTS}}": "Principal's comments"
            },
            "Behavioral Assessment": {
                "{{ATA}}": "Attitude to Authority",
                "{{RESP}}": "Responsibility",
                "{{CO_OP}}": "Cooperation",
                "{{ATC}}": "Attitude to Corrections",
                "{{LEAD}}": "Leadership",
                "{{DEP}}": "Deportment",
                "{{SOC}}": "Sociability",
                "{{INIT}}": "Initiative",
                "{{CON_MGT}}": "Conflict Management",
                "{{APP}}": "Application"
            }
        }
        
        for category, placeholders in placeholder_data.items():
            st.write(f"**{category}**")
            for placeholder, description in placeholders.items():
                st.write(f"- `{placeholder}` - {description}")
