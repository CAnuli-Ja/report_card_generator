import pandas as pd
from docx import Document
import os
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk

class ReportCardGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("JWFCSS Report Card Generator")
        self.root.geometry("700x650")
        self.root.resizable(False, True)
        
        # Store file paths
        self.excel_path = tk.StringVar()
        self.template_path = tk.StringVar()
        self.output_dir = tk.StringVar()
        self.sheet_name = tk.StringVar()
        
        # Create main container
        self.container = tk.Frame(root)
        self.container.pack(side="top", fill="both", expand=True)
        self.container.grid_rowconfigure(0, weight=1)
        self.container.grid_columnconfigure(0, weight=1)
        
        # Create navigation bar
        self.create_navbar()
        
        # Create frames for different pages
        self.frames = {}
        for F in (LandingPage, UploadPage):
            frame = F(self.container, self)
            self.frames[F] = frame
            frame.grid(row=1, column=0, sticky="nsew")
        
        # Show landing page first
        self.show_frame(LandingPage)
    
    def create_navbar(self):
        """Create navigation bar at top"""
        navbar = tk.Frame(self.root, bg="#2c3e50", height=60)
        navbar.pack(side="top", fill="x")
        navbar.pack_propagate(False)
        
        # Logo/Title
        title_label = tk.Label(
            navbar,
            text="JWFCSS Report Card Generator",
            bg="#2c3e50",
            fg="white",
            font=("Arial", 14, "bold")
        )
        title_label.pack(side="left", padx=20, pady=10)
        
        # Navigation buttons
        nav_frame = tk.Frame(navbar, bg="#2c3e50")
        nav_frame.pack(side="right", padx=20, pady=10)
        
        home_btn = tk.Button(
            nav_frame,
            text="Home",
            bg="#3498db",
            fg="white",
            relief="flat",
            padx=15,
            command=lambda: self.show_frame(LandingPage)
        )
        home_btn.pack(side="left", padx=5)
        
        upload_btn = tk.Button(
            nav_frame,
            text="Generate Reports",
            bg="#27ae60",
            fg="white",
            relief="flat",
            padx=15,
            command=lambda: self.show_frame(UploadPage)
        )
        upload_btn.pack(side="left", padx=5)
    
    def show_frame(self, cont):
        """Display a frame"""
        frame = self.frames[cont]
        frame.tkraise()


class LandingPage(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.controller = controller
        
        # Main content frame
        main_frame = tk.Frame(self)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=30, pady=30)
        
        # Title
        title = tk.Label(
            main_frame,
            text="Welcome to Report Card Generator",
            font=("Arial", 20, "bold"),
            fg="#2c3e50"
        )
        title.pack(pady=20)
        
        # Description
        description = tk.Label(
            main_frame,
            text="Automatically generate personalized report cards for all your students\nusing Excel gradesheets and Word templates.",
            font=("Arial", 12),
            justify=tk.CENTER,
            fg="#34495e"
        )
        description.pack(pady=10)
        
        # Features section
        features_frame = tk.LabelFrame(
            main_frame,
            text="Features",
            font=("Arial", 12, "bold"),
            padx=20,
            pady=20
        )
        features_frame.pack(fill=tk.BOTH, expand=True, pady=20)
        
        features = [
            "✓ Batch generate report cards for entire classes",
            "✓ Automatic grade-based comments",
            "✓ Support for multiple subjects and behavioral assessments",
            "✓ Customizable Word templates",
            "✓ Process multiple classes in one workbook",
            "✓ Fast and efficient document creation"
        ]
        
        for feature in features:
            feature_label = tk.Label(
                features_frame,
                text=feature,
                font=("Arial", 11),
                justify=tk.LEFT,
                fg="#2c3e50"
            )
            feature_label.pack(anchor="w", pady=5)
        
        # Quick start section
        quickstart_frame = tk.LabelFrame(
            main_frame,
            text="Quick Start",
            font=("Arial", 12, "bold"),
            padx=20,
            pady=15
        )
        quickstart_frame.pack(fill=tk.BOTH, pady=20)
        
        steps = [
            "1. Prepare your Excel gradesheet with student data",
            "2. Create a Word template with placeholders",
            "3. Click 'Generate Reports' and select your files",
            "4. Your report cards will be created in seconds!"
        ]
        
        for step in steps:
            step_label = tk.Label(
                quickstart_frame,
                text=step,
                font=("Arial", 10),
                justify=tk.LEFT,
                fg="#2c3e50"
            )
            step_label.pack(anchor="w", pady=3)
        
        # Get started button
        get_started_btn = tk.Button(
            main_frame,
            text="Get Started →",
            font=("Arial", 12, "bold"),
            bg="#27ae60",
            fg="white",
            relief="flat",
            padx=30,
            pady=10,
            command=lambda: self.controller.show_frame(UploadPage)
        )
        get_started_btn.pack(pady=20)


class UploadPage(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self, parent)
        self.controller = controller
        
        # Main content frame
        main_frame = tk.Frame(self)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Excel Gradesheet section
        tk.Label(main_frame, text="Excel Gradesheet", font=("Arial", 10, "bold")).pack(anchor="w", pady=(10, 2))
        tk.Entry(main_frame, textvariable=controller.excel_path, width=80).pack()
        tk.Button(main_frame, text="Browse Excel Gradesheet", command=self.browse_excel, width=40).pack(pady=5)
        
        # Sheet Name section
        tk.Label(main_frame, text="Sheet Name", font=("Arial", 10, "bold")).pack(anchor="w", pady=(10, 2))
        self.sheet_dropdown = ttk.Combobox(main_frame, textvariable=controller.sheet_name, width=77, state="readonly")
        self.sheet_dropdown.pack(fill=tk.X, padx=2, pady=5)
        
        # Word Report Template section
        tk.Label(main_frame, text="Word Report Template", font=("Arial", 10, "bold")).pack(anchor="w", pady=(10, 2))
        tk.Entry(main_frame, textvariable=controller.template_path, width=80).pack()
        tk.Button(main_frame, text="Browse Word Report Template", command=self.browse_template, width=40).pack(pady=5)
        
        # Output Folder section
        tk.Label(main_frame, text="Output Folder", font=("Arial", 10, "bold")).pack(anchor="w", pady=(10, 2))
        tk.Entry(main_frame, textvariable=controller.output_dir, width=80).pack()
        tk.Button(main_frame, text="Browse Output Folder", command=self.browse_output, width=40).pack(pady=5)
        
        # Generate button
        tk.Button(
            main_frame,
            text="Generate Report Cards",
            command=self.generate_reports,
            height=2,
            width=30,
            font=("Arial", 12, "bold"),
            bg="#4CAF50",
            fg="white"
        ).pack(pady=15)
    
    def browse_excel(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx")])
        if file_path:
            print(f"Selected Excel file: {file_path}")
            self.controller.excel_path.set(file_path)
            self.load_sheet_names(file_path)
    
    def load_sheet_names(self, file_path):
        """Load sheet names from the selected Excel file"""
        try:
            xls = pd.ExcelFile(file_path)
            sheets = xls.sheet_names
            print(f"Found sheets: {sheets}")
            self.sheet_dropdown['values'] = sheets
            if sheets:
                self.controller.sheet_name.set(sheets[0])
                print(f"Set default sheet to: {sheets[0]}")
        except Exception as e:
            print(f"Error loading sheets: {str(e)}")
            messagebox.showerror("Error", f"Could not read sheet names: {str(e)}")
    
    def browse_template(self):
        self.controller.template_path.set(filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")]))
    
    def browse_output(self):
        self.controller.output_dir.set(filedialog.askdirectory())
    
    def generate_reports(self):
        if not self.controller.excel_path.get() or not self.controller.template_path.get() or not self.controller.output_dir.get() or not self.controller.sheet_name.get():
            messagebox.showerror("Missing Information", "Please select all required files, output folder, and sheet name.")
            return

        try:
            # Read Excel file with specified sheet name
            raw = pd.read_excel(self.controller.excel_path.get(), sheet_name=self.controller.sheet_name.get(), header=None)
            
            # Find the row where actual data starts
            start_row = None
            for idx, val in enumerate(raw[0]):
                if pd.notna(val) and str(val).strip().isdigit():
                    start_row = idx
                    break
            
            if start_row is None:
                messagebox.showerror("Error", "Could not find student data in Excel file.")
                return
            
            # Extract data starting from the header row
            df = raw.iloc[start_row:].reset_index(drop=True)
            df.columns = df.iloc[0]
            df = df[1:].reset_index(drop=True)
            
            if len(df) == 0:
                messagebox.showerror("Error", "No student data found in Excel file.")
                return

            os.makedirs(self.controller.output_dir.get(), exist_ok=True)
            
            count = 0
            for _, row in df.iterrows():
                doc = Document(self.controller.template_path.get())

                # Get student data
                first_name = str(row.iloc[1]).strip() if pd.notna(row.iloc[1]) else ""
                last_name = str(row.iloc[2]).strip() if pd.notna(row.iloc[2]) else ""
                full_name = f"{first_name} {last_name}".strip()
                
                if not full_name:
                    continue

                # Extract all grades and data
                replacements = {}
                subject_data = [
                    ("{{HFL}}", 3), ("{{LIT}}", 4), ("{{REL}}", 5), ("{{ELANG}}", 6),
                    ("{{ELIT}}", 7), ("{{FR}}", 8), ("{{SPN}}", 9), ("{{SSTUD}}", 10),
                    ("{{HIST}}", 11), ("{{MATH}}", 12), ("{{AGR}}", 13), ("{{INTSCI}}", 14),
                    ("{{PE}}", 15), ("{{IT}}", 16), ("{{VA}}", 17), ("{{HE}}", 18),
                    ("{{TOTAL}}", 19), ("{{SAVG}}", 20), ("{{TOTAL_SUBJECTS}}", 21),
                    ("{{PASSED}}", 22), ("{{FORM_TEACHER_COMMENTS}}", 23),
                    ("{{PRINCIPAL_COMMENTS}}", 24), ("{{ATA}}", 25), ("{{RESP}}", 26),
                    ("{{CO_OP}}", 27), ("{{ATC}}", 28), ("{{LEAD}}", 29), ("{{DEP}}", 30),
                    ("{{SOC}}", 31), ("{{INIT}}", 32), ("{{CON_MGT}}", 33), ("{{APP}}", 34)
                ]
                
                replacements["{{NAME}}"] = full_name
                
                for placeholder, col_idx in subject_data:
                    value = str(row.iloc[col_idx]).strip() if pd.notna(row.iloc[col_idx]) else ""
                    replacements[placeholder] = value
                
                # Helper function to replace text
                def replace_text_in_paragraph(paragraph, replacements):
                    for placeholder, value in replacements.items():
                        if placeholder in paragraph.text:
                            full_text = paragraph.text
                            new_text = full_text.replace(placeholder, value)
                            for run in paragraph.runs:
                                run.text = ""
                            if paragraph.runs:
                                paragraph.runs[0].text = new_text
                            else:
                                paragraph.add_run(new_text)

                # Replace in paragraphs
                for p in doc.paragraphs:
                    replace_text_in_paragraph(p, replacements)

                # Replace in tables
                for table in doc.tables:
                    for row_table in table.rows:
                        for cell in row_table.cells:
                            for p in cell.paragraphs:
                                replace_text_in_paragraph(p, replacements)

                filename = f"{full_name.replace(' ', '_')}_Report.docx"
                doc.save(os.path.join(self.controller.output_dir.get(), filename))
                count += 1

            messagebox.showinfo("Success", f"{count} report card(s) generated successfully!")

        except Exception as e:
            import traceback
            error_msg = f"{str(e)}\n\n{traceback.format_exc()}"
            messagebox.showerror("Error", error_msg)


if __name__ == "__main__":
    root = tk.Tk()
    app = ReportCardGeneratorApp(root)
    root.mainloop()
