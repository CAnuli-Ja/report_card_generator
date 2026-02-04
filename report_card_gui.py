import pandas as pd
from docx import Document
import os
import tkinter as tk
from tkinter import filedialog, messagebox

# ---------------- AUTO COMMENT ---------------- #
def auto_comment(avg):
    try:
        avg = float(avg)
        if avg >= 75:
            return "Excellent performance. The student demonstrates strong understanding and consistent effort."
        elif avg >= 60:
            return "Good performance. Continued effort and focus will yield further improvement."
        elif avg >= 50:
            return "Satisfactory performance. Greater consistency and engagement are encouraged."
        else:
            return "Performance needs improvement. Increased effort and academic support are recommended."
    except:
        return "Performance reviewed for the term."

# ---------------- MAIN LOGIC ---------------- #
def generate_reports():
    if not excel_path.get() or not template_path.get() or not output_dir.get():
        messagebox.showerror("Missing Information", "Please select all required files and output folder.")
        return

    try:
        # Read Excel file
        raw = pd.read_excel(excel_path.get(), header=None)
        # Find the row where actual data starts (first row with a number in column 0)
        start_row = None
        for idx, val in enumerate(raw[0]):
            if pd.notna(val) and str(val).strip().isdigit():
                start_row = idx
                break
        
        if start_row is None:
            messagebox.showerror("Error", "Could not find student data in Excel file. Make sure the first column contains student numbers.")
            return
        
        # Extract data starting from the header row
        df = raw.iloc[start_row:].reset_index(drop=True)
        df.columns = df.iloc[0]
        df = df[1:].reset_index(drop=True) # Remove header row from data
        
        # Display dataframe in console
        #print("\n--- Student Data ---")
        #print(df)
        #print("-------------------\n")
        
        
        if len(df) == 0:
            messagebox.showerror("Error", "No student data found in Excel file.")
            return

        os.makedirs(output_dir.get(), exist_ok=True)
        
        count = 0
        for _, row in df.iterrows():
            doc = Document(template_path.get())

            # Get student name (assuming columns 1 and 2 are first and last name)
            first_name = str(row.iloc[1]).strip() if pd.notna(row.iloc[1]) else ""
            last_name = str(row.iloc[2]).strip() if pd.notna(row.iloc[2]) else ""
            full_name = f"{first_name} {last_name}".strip()
            health_family_life = str(row.iloc[3]).strip() if pd.notna(row.iloc[3]) else ""
            literacy = str(row.iloc[4]).strip() if pd.notna(row.iloc[4]) else ""
            religion = str(row.iloc[5]).strip() if pd.notna(row.iloc[5]) else ""
            english_language = str(row.iloc[6]).strip() if pd.notna(row.iloc[6]) else ""
            english_literature = str(row.iloc[7]).strip() if pd.notna(row.iloc[7]) else ""
            french = str(row.iloc[8]).strip() if pd.notna(row.iloc[8]) else ""
            spanish = str(row.iloc[9]).strip() if pd.notna(row.iloc[9]) else ""
            social_studies = str(row.iloc[10]).strip() if pd.notna(row.iloc[10]) else ""
            history = str(row.iloc[11]).strip() if pd.notna(row.iloc[11]) else ""
            mathematics = str(row.iloc[12]).strip() if pd.notna(row.iloc[12]) else ""
            agricultural_science = str(row.iloc[13]).strip() if pd.notna(row.iloc[13]) else ""
            integrated_science = str(row.iloc[14]).strip() if pd.notna(row.iloc[14]) else ""
            physical_education = str(row.iloc[15]).strip() if pd.notna(row.iloc[15]) else ""
            information_technology = str(row.iloc[16]).strip() if pd.notna(row.iloc[16]) else ""
            visual_art = str(row.iloc[17]).strip() if pd.notna(row.iloc[17]) else ""
            home_economics = str(row.iloc[18]).strip() if pd.notna(row.iloc[18]) else ""
            total = str(row.iloc[19]).strip() if pd.notna(row.iloc[19]) else ""
            average = str(row.iloc[20]).strip() if pd.notna(row.iloc[20]) else ""
            total_subjects = str(row.iloc[21]).strip() if pd.notna(row.iloc[21]) else ""
            subject_passed = str(row.iloc[22]).strip() if pd.notna(row.iloc[22]) else ""
            form_teacher_comments = str(row.iloc[23]).strip() if pd.notna(row.iloc[23]) else ""
            principal_comments = str(row.iloc[24]).strip() if pd.notna(row.iloc[24]) else ""
            attitude_to_authority = str(row.iloc[25]).strip() if pd.notna(row.iloc[25]) else ""
            responsibility = str(row.iloc[26]).strip() if pd.notna(row.iloc[26]) else ""
            cooperation = str(row.iloc[27]).strip() if pd.notna(row.iloc[27]) else ""
            attitude_to_corrections = str(row.iloc[28]).strip() if pd.notna(row.iloc[28]) else ""
            leadership = str(row.iloc[29]).strip() if pd.notna(row.iloc[29]) else ""
            deportment = str(row.iloc[30]).strip() if pd.notna(row.iloc[30]) else ""
            sociability = str(row.iloc[31]).strip() if pd.notna(row.iloc[31]) else ""
            initiative = str(row.iloc[32]).strip() if pd.notna(row.iloc[32]) else ""
            conflict_management = str(row.iloc[33]).strip() if pd.notna(row.iloc[33]) else ""
            application = str(row.iloc[34]).strip() if pd.notna(row.iloc[34]) else ""
            
            # Display full name in console for visibility
            #print(f"\nGenerating report for: {full_name}")
            
            # Debug: Show all text found in template
            all_text = [p.text for p in doc.paragraphs]
            all_table_text = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            all_table_text.append(p.text)
            
            #print(f"  Document text found: {[t for t in all_text if t.strip()][:5]}")  # Show first 5 non-empty paragraphs
            #print(f"  Table text found: {[t for t in all_table_text if t.strip()][:5]}")  # Show first 5 non-empty table cells
            
            if not full_name:
                continue

            # Define replacements mapping with {{placeholder}} format
            replacements = {
                "{{NAME}}": full_name,
                "{{HFL}}": health_family_life,
                "{{LIT}}": literacy,
                "{{REL}}": religion,
                "{{ELANG}}": english_language,
                "{{ELIT}}": english_literature,
                "{{FR}}": french,
                "{{SPN}}": spanish,
                "{{SSTUD}}": social_studies,
                "{{HIST}}": history,
                "{{MATH}}": mathematics,
                "{{AGR}}": agricultural_science,
                "{{INTSCI}}": integrated_science,
                "{{PE}}": physical_education,
                "{{IT}}": information_technology,
                "{{VA}}": visual_art,
                "{{HE}}": home_economics,
                "{{TOTAL}}": total,
                "{{SAVG}}": average,
                "{{TOTAL_SUBJECTS}}": total_subjects,
                "{{PASSED}}": subject_passed,
                "{{FORM_TEACHER_COMMENTS}}": form_teacher_comments,
                "{{PRINCIPAL_COMMENTS}}": principal_comments,
                "{{ATA}}": attitude_to_authority,
                "{{RESP}}": responsibility,
                "{{CO_OP}}": cooperation,
                "{{ATC}}": attitude_to_corrections,
                "{{LEAD}}": leadership,
                "{{DEP}}": deportment,
                "{{SOC}}": sociability,
                "{{INIT}}": initiative,
                "{{CON_MGT}}": conflict_management,
                "{{APP}}": application,
            }

            # Helper function to replace text in paragraph (handles runs split across formatting)
            def replace_text_in_paragraph(paragraph, replacements):
                for placeholder, value in replacements.items():
                    if placeholder in paragraph.text:
                        # Combine all runs text
                        full_text = paragraph.text
                        if placeholder in full_text:
                            # Remove placeholder if value is empty, otherwise replace with value
                            if value:
                                new_text = full_text.replace(placeholder, value)
                            else:
                                new_text = full_text.replace(placeholder, "")
                            # Clear all runs and add new text to first run
                            for run in paragraph.runs:
                                run.text = ""
                            if paragraph.runs:
                                paragraph.runs[0].text = new_text
                            else:
                                paragraph.add_run(new_text)

            # Replace placeholders in paragraphs
            for p in doc.paragraphs:
                replace_text_in_paragraph(p, replacements)

            # Replace placeholders in tables
            for table in doc.tables:
                for row_table in table.rows:
                    for cell in row_table.cells:
                        for p in cell.paragraphs:
                            replace_text_in_paragraph(p, replacements)

            filename = f"{full_name.replace(' ', '_')}_Report.docx"
            doc.save(os.path.join(output_dir.get(), filename))
            count += 1

        messagebox.showinfo("Success", f"{count} report card(s) generated successfully!")

    except Exception as e:
        import traceback
        error_msg = f"{str(e)}\n\n{traceback.format_exc()}"
        messagebox.showerror("Error", error_msg)

# ---------------- FILE PICKERS ---------------- #
def browse_excel():
    excel_path.set(filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx")]))

def browse_template():
    template_path.set(filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")]))

def browse_output():
    output_dir.set(filedialog.askdirectory())

# ---------------- GUI SETUP ---------------- #
root = tk.Tk()
root.title("JWFCSS Report Card Generator")
root.geometry("580x420")
root.resizable(False, False)

excel_path = tk.StringVar()
template_path = tk.StringVar()
output_dir = tk.StringVar()

tk.Label(root, text="Excel Gradesheet").pack(pady=5)
tk.Entry(root, textvariable=excel_path, width=70).pack()
tk.Button(root, text="Browse Excel Gradesheet", command=browse_excel).pack(pady=5)

tk.Label(root, text="Word Report Template").pack(pady=5)
tk.Entry(root, textvariable=template_path, width=70).pack()
tk.Button(root, text="Browse Word Report Template", command=browse_template).pack(pady=5)

tk.Label(root, text="Output Folder").pack(pady=5)
tk.Entry(root, textvariable=output_dir, width=70).pack()
tk.Button(root, text="Browse Output Folder", command=browse_output).pack(pady=5)
tk.Button(
    root,
    text="Generate Report Cards",
    command=generate_reports,
    height=2,
    width=25,
    font=("Arial", 12, "bold")
).pack(pady=20)

root.mainloop()
